import re
import docx
import os
import yaml
from docx import Document
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml import OxmlElement

def get_qn_name(tag_name):
    return qn(tag_name)

def get_numbering_part(doc):
    numbering_part_element = doc.part.numbering_part.element
    w_val = get_qn_name('w:val')
    # 1. Locate all w:num elements. Each w:num element contains a numId attribute that associates it with a paragraph.
    w_num = get_qn_name('w:num')
    num_elements = numbering_part_element.findall(w_num)
    
    # 2. Locate all abstractNumId elements. Typically, there is one abstractNumId element.
    absNumId_to_numId = {}
    w_abstractNumId = get_qn_name('w:abstractNumId')
    w_numId = get_qn_name('w:numId')
    for num_element in num_elements:
        abstractNumId = num_element.findall(w_abstractNumId)
        if len(abstractNumId) == 0:
            continue
        abstractNumId = abstractNumId[0]
        abstractNumId = abstractNumId.get(w_val)
        numId = num_element.get(w_numId)
        if abstractNumId is not None and numId is not None:
            absNumId_to_numId.update({abstractNumId: numId})
    
    # 3. Locate all abstractNum elements.
    w_abstractNum = get_qn_name('w:abstractNum')
    abstractNum_elements = numbering_part_element.findall(w_abstractNum)
    
    # 4. Within each abstractNum element, examine the abstractNumId, lvl, lvlText, and numFmt elements. 
    # Under normal circumstances, there would be only one lvlText and one numFmt element for each level.
    w_lvl = get_qn_name('w:lvl')
    w_ilvl = get_qn_name('w:ilvl')
    w_lvlText = get_qn_name('w:lvlText')
    w_numFmt = get_qn_name('w:numFmt')
    
    numbering_part = {}
    for abstractNum_element in abstractNum_elements:
        abstractNumId = abstractNum_element.get(w_abstractNumId)
        if abstractNumId is None:
            continue
        bucket = {}
        lvl_elements = abstractNum_element.findall(w_lvl)
        numFmt = 'decimal'
        for lvl_element in lvl_elements:
            ilvl = lvl_element.get(w_ilvl)
            if ilvl is None:
                continue
                
            lvlText_elements = lvl_element.findall(w_lvlText)
            numFmt_elements = lvl_element.findall(w_numFmt)
            if len(lvlText_elements) == 0 or len(numFmt_elements) == 0:
                continue
            numFmt = numFmt_elements[0].get(w_val)
            text = lvlText_elements[0].get(w_val)
            bucket.update({ilvl:[text, numFmt]})
            
        if abstractNumId in absNumId_to_numId.keys():
            numbering_part.update({
                absNumId_to_numId[abstractNumId]: bucket
            })
    return numbering_part

def get_known_formats():
    def generate_roman_numerals(limit = 100, uppercase = True):
        roman_map = { 1: 'I', 4: 'IV', 5: 'V', 9: 'IX', 10: 'X', 40: 'XL', 50: 'L', 90: 'XC', 100: 'C'}

        roman_numerals = []
        for i in range(1, limit + 1):
            result = ""
            for value, numeral in sorted(roman_map.items(), reverse=True):
                while i >= value:
                    result += numeral
                    i -= value
            result = result if uppercase else result.lower()
            roman_numerals.append(result)
        return roman_numerals
    
    ENG_LETTERS = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
    upperLetter = list(ENG_LETTERS) + [
        f'{a}{b}'
        for a in list(ENG_LETTERS) for b in list(ENG_LETTERS)   
    ]
    lowerLetter = list(ENG_LETTERS.lower()) + [
        f'{a}{b}'
        for a in list(ENG_LETTERS.lower()) for b in list(ENG_LETTERS.lower())
    ]
    upperRoman = generate_roman_numerals()
    lowerRoman = generate_roman_numerals(uppercase = False)
    return {
        'upperLetter': upperLetter,
        'lowerLetter': lowerLetter,
        'upperRoman': upperRoman,
        'lowerRoman': lowerRoman
    }

def get_string_for_format(format, stack_number, known_formats):
    if format in known_formats.keys():
        if len(known_formats[format]) > stack_number:
            return known_formats[format][stack_number]
    return stack_number+1

def apply_numbering(numId_val, ilvl_val, numbering_part, numbering_part_stack, known_formats):
    numbering = numbering_part[numId_val]
    if not numId_val in numbering_part_stack.keys():
        numbering_part_stack.update({numId_val: {}})
        
    number_format, format = numbering[ilvl_val]
    ilvl_val = int(ilvl_val)
    if not ilvl_val in numbering_part_stack[numId_val].keys():
        numbering_part_stack[numId_val].update({ilvl_val: 0})
    else:
        for drop in range(max(numbering_part_stack[numId_val].keys())-ilvl_val):
            drop = ilvl_val+drop+1
            if drop in numbering_part_stack[numId_val].keys():
                numbering_part_stack[numId_val].pop(drop)
        numbering_part_stack[numId_val][ilvl_val] += 1
    search = re.findall(r'(\%\d+)', number_format)
    if len(search) == 0:
        return number_format
    number_format = re.sub(r'(\%\d+)', '{}', number_format)
    
    format_letters = []
    for _ in range(len(search)):
        stack_number = numbering_part_stack[numId_val][ilvl_val]
        _, format = numbering[str(ilvl_val)]
        letter = get_string_for_format(format, stack_number, known_formats)
        format_letters.append(letter)
        ilvl_val -= 1
    
    return number_format.format(*format_letters[::-1]), numbering_part_stack

def get_ppr_val(para):
    if not isinstance(para, docx.oxml.text.paragraph.CT_P):
        return None, None
    ilvl_val = None
    numId_val = None
    val_name = docx.oxml.ns.qn('w:val')
    ppr = para.pPr
    if ppr is not None:
        numpr = ppr.numPr
        if numpr is not None:
            numId = numpr.numId
            if numId is not None:
                numId_val = numId.get(val_name)    
            ilvl = numpr.ilvl
            if ilvl is not None:
                ilvl_val = ilvl.get(val_name)
            
    return numId_val, ilvl_val

def add_numbered_paragraph(doc, text, level):
    # 1. 创建段落对象
    p = doc.add_paragraph()

    run = p.add_run(text)
    
    # 2. 创建编号属性
    numPr = OxmlElement('w:numPr')
    
    # 设置缩进级别
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    
    # 3. 根据级别设置不同编号样式
    numId = OxmlElement('w:numId')
    if level == 0:  # 一级标题
        numId.set(qn('w:val'), '7')  # 阿拉伯数字（1, 1）[7](@ref)
    elif level == 1:  # 二级标题
        numId.set(qn('w:val'), '7')   # 阿拉伯数字（1.1, 1.2）[7](@ref)

    numPr.append(ilvl)
    numPr.append(numId)

    # 4. 将属性注入段落
    pPr = p._element.get_or_add_pPr()
    pPr.append(numPr)
    
    # 5. 根据级别设置差异化缩进
    ind = OxmlElement('w:ind')
    if level == 0:
        ind.set(qn('w:left'), "425")       # 一级不缩进
        ind.set(qn('w:leftChars'), "0")  # 悬挂缩进
        ind.set(qn('w:hanging'), "425")  # 悬挂缩进
        ind.set(qn('w:firstLineChars'), "0")  # 悬挂缩进
    elif level == 1:
        ind.set(qn('w:left'), "567")       # 一级不缩进
        ind.set(qn('w:leftChars'), "0")  # 悬挂缩进
        ind.set(qn('w:hanging'), "567")  # 悬挂缩进
        ind.set(qn('w:firstLineChars'), "0")  # 悬挂缩进
    pPr.append(ind)

def add_numbered_head(doc, text, level):
    # 1. 创建段落对象
    p = doc.add_heading()

    run = p.add_run(text)
    
    # 2. 创建编号属性
    numPr = OxmlElement('w:numPr')
    
    # 设置缩进级别
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    
    # 3. 根据级别设置不同编号样式
    numId = OxmlElement('w:numId')
    if level == 0:  # 一级标题
        numId.set(qn('w:val'), '7')  # 阿拉伯数字（1, 1）[7](@ref)
    elif level == 1:  # 二级标题
        numId.set(qn('w:val'), '7')   # 阿拉伯数字（1.1, 1.2）[7](@ref)

    numPr.append(ilvl)
    numPr.append(numId)

    # 4. 将属性注入段落
    pPr = p._element.get_or_add_pPr()
    pPr.append(numPr)
    
    # 5. 根据级别设置差异化缩进
    ind = OxmlElement('w:ind')
    if level == 0:
        ind.set(qn('w:left'), "425")       # 一级不缩进
        ind.set(qn('w:leftChars'), "0")  # 悬挂缩进
        ind.set(qn('w:hanging'), "425")  # 悬挂缩进
        ind.set(qn('w:firstLineChars'), "0")  # 悬挂缩进
    elif level == 1:
        ind.set(qn('w:left'), "567")       # 一级不缩进
        ind.set(qn('w:leftChars'), "0")  # 悬挂缩进
        ind.set(qn('w:hanging'), "567")  # 悬挂缩进
        ind.set(qn('w:firstLineChars'), "0")  # 悬挂缩进
    pPr.append(ind)
    
    # 6. 设置中文字体（防乱码）
    # run.font.name = '宋体'
    # run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 关键设置[3](@ref)

def load_custom_numbering(doc, numbering_xml_path):
    """加载自定义编号XML到文档"""
    # 读取保存的编号XML文件
    with open(numbering_xml_path, 'r', encoding='utf-8') as f:
        numbering_xml = f.read()
    
    # 解析XML内容[1,9](@ref)
    new_numbering = parse_xml(numbering_xml)
    
    # 获取文档的编号部件
    numbering_part = doc.part.numbering_part
    
    # 清空现有编号定义
    for elem in numbering_part._element.getchildren():
        numbering_part._element.remove(elem)
    
    # 添加自定义编号定义[1,9](@ref)
    for child in list(new_numbering):
        numbering_part._element.append(child)
    
    return numbering_part

def create_document_with_custom_numbering(numbering_xml_path):
    """创建带有自定义编号的新文档"""
    # 创建空白文档
    doc = Document()
    
    # 加载自定义编号
    load_custom_numbering(doc, numbering_xml_path)
    
    # 创建关联自定义编号的样式[1,2](@ref)
    style_id = "CustomNumberingStyle"
    if style_id not in doc.styles:
        style = doc.styles.add_style(style_id, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        pPr = style._element.get_or_add_pPr()
        numPr = pPr.get_or_add_numPr()
        numId = numPr.get_or_add_numId()
        numId.val = 1  # 对应XML中的numId
    
    return doc

def create_document_with_programmatic_numbering(config_path: str = './config/docx_config.yaml'):
    """创建带有程序化定义的多级编号的新文档（无需外部XML）。
    优先按配置文件 headings.h2/h3/... 下的 numbering.format 生成 lvlText（如 "%1"、"%1.%2"）。
    """
    doc = Document()
    numbering_part = doc.part.numbering_part

    # 根节点
    numbering_root = numbering_part._element

    # 创建 abstractNum（抽象编号定义），id=0（对齐 XML）
    abstract_num = OxmlElement('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), '0')
    # 多级类型：multilevel（对齐 XML）
    multi_type = OxmlElement('w:multiLevelType')
    multi_type.set(qn('w:val'), 'multilevel')
    abstract_num.append(multi_type)

    # 读取配置中的编号格式：h2 对应 ilvl=0，h3 -> ilvl=1 ...
    level_formats = {}
    try:
        if os.path.exists(config_path):
            with open(config_path, 'r', encoding='utf-8') as f:
                cfg = yaml.safe_load(f) or {}
            headings = (cfg.get('headings') or {})
            for lvl_name, idx in [('h2', 0), ('h3', 1), ('h4', 2), ('h5', 3), ('h6', 4), ('h7', 5)]:
                heading_cfg = headings.get(lvl_name) or {}
                numbering = heading_cfg.get('numbering') or {}
                fmt = numbering.get('format')
                if fmt:
                    level_formats[idx] = fmt
    except Exception:
        pass

    # 定义前6级（0-5）十进制多级编号。
    # 为了让编号与标题之间仅保留一个空格，这里不再设置段落缩进/悬挂，
    # 而是使用 w:suff="space" 实现紧凑分隔。
    for ilvl in range(6):
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(ilvl))

        # 格式
        num_fmt = OxmlElement('w:numFmt')
        num_fmt.set(qn('w:val'), 'decimal')
        lvl.append(num_fmt)

        # 文本模板：优先使用配置的 format，否则默认 "%1.%2."（带结尾点，对齐 XML）
        if ilvl in level_formats:
            # 若配置未带结尾点，则补一个点
            fmt_val = str(level_formats[ilvl])
            lvl_text_val = fmt_val if fmt_val.endswith('.') else fmt_val + '.'
        else:
            parts = [f"%{i+1}" for i in range(ilvl+1)]
            lvl_text_val = '.'.join(parts) + '.'
        lvl_text = OxmlElement('w:lvlText')
        lvl_text.set(qn('w:val'), lvl_text_val)
        lvl.append(lvl_text)

        # 起始值 1
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        # 对齐方式
        lvl_jc = OxmlElement('w:lvlJc')
        lvl_jc.set(qn('w:val'), 'left')
        lvl.append(lvl_jc)

        # 后缀使用一个空格，且不引入额外缩进/悬挂
        suff = OxmlElement('w:suff')
        suff.set(qn('w:val'), 'space')
        lvl.append(suff)

        abstract_num.append(lvl)

    numbering_root.append(abstract_num)

    # 创建 num（实例），numId=7，关联 abstractNumId=0（对齐 XML）
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), '7')
    abstract_num_id = OxmlElement('w:abstractNumId')
    abstract_num_id.set(qn('w:val'), '0')
    num.append(abstract_num_id)
    numbering_root.append(num)

    # 创建样式但不强绑 numPr，避免与段落注入冲突
    style_id = "ProgrammaticNumberingStyle"
    if style_id not in doc.styles:
        style = doc.styles.add_style(style_id, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]

    return doc

# 测试代码示例
if __name__ == "__main__":
    try:
        doc = Document("numbered1.docx")  # 请替换为实际的Word文档路径
        numbering_part = get_numbering_part(doc)
        print("提取的编号格式配置:")
        print(numbering_part)
    except Exception as e:
        print(f"处理文档时出错: {e}")
        print("请确保提供了有效的Word文档路径")

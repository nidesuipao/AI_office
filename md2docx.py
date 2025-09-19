from ctypes import alignment
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

import markdown
from markdown import Markdown
from markdown.treeprocessors import Treeprocessor
import yaml
from bs4 import BeautifulSoup, Tag
from bs4.element import NavigableString
from typing import Optional
import os

from utils.docx_utils import *

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from markdown import Markdown
from bs4 import BeautifulSoup, Tag
from bs4.element import NavigableString
from latex2mathml.converter import convert as latex_to_mathml
from lxml import etree

class MarkdownToDocxConverter:
    def __init__(self, numbering_config='./config/docx_numbering.xml', style_config=None):
        self.numbering_config = numbering_config
        self.style_config = self.load_config()
        self.counters = [0] * 6
        self._xslt_transform = None

    def load_config(self, config_path="/home/yzy/document/project/AI_office-main/config/docx_config.yaml"):
        with open(config_path, 'r', encoding='utf-8') as f:
            config = yaml.safe_load(f)
        
        def process_config_section(section):
            if isinstance(section, dict):
                for key, value in section.items():
                    if key == 'size':
                        section[key] = Pt(value)
                    elif key == 'color' and isinstance(value, str):
                        hex_color = value.lstrip('#')
                        section[key] = RGBColor(*tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4)))
                    elif key in ['space_before', 'space_after', 'line_spacing']:
                        # 处理段落间距相关配置
                        if key == 'line_spacing':
                            section[key] = value  # 行间距保持原值
                        else:
                            section[key] = Pt(value)  # 前后间距转换为Pt
                    elif key in ['top', 'bottom', 'left', 'right'] and 'margins' in str(section):
                        # 处理页面边距，保持原始数值，在应用时转换为Inches
                        pass
                    elif key == 'unit' and value == 'inches':
                        # 单位转换标记，保持原值
                        pass
                    elif isinstance(value, (dict, list)):
                        process_config_section(value)

            elif isinstance(section, list):
                for item in section:
                    process_config_section(item)
        
        process_config_section(config)
        return config

    
    def _default_style_config(self):
        """默认样式配置"""
        return {
            'headings': {
                'h1': {'font': {'western': 'Arial', 'east_asian': '黑体', 'size': Pt(22), 'color': RGBColor(0, 0, 0)}},
                'h2': {'font': {'western': 'Arial', 'east_asian': '黑体', 'size': Pt(18), 'color': RGBColor(0, 0, 0)}},
                'h3': {'font': {'western': 'Arial', 'east_asian': '黑体', 'size': Pt(18), 'color': RGBColor(0, 0, 0)}},
                # 其他标题样式...
            },
            'paragraph': {
                'font': {'western': 'Times New Roman', 'east_asian': '宋体', 'size': Pt(14), 'color': RGBColor(0, 0, 0)}
            }
        }
    
    def _create_document(self):
        """创建带有自定义编号的文档"""
        # 返回普通空白文档；后续采用前缀文本编号（无关联）
        return Document()
    
    def _apply_page_layout(self):
        """应用页面布局设置"""
        page_layout = self.style_config.get('page_layout', {})
        if page_layout:
            # 设置页面边距
            margins = page_layout.get('margins', {})
            if margins:
                self.doc.sections[0].top_margin = Inches(margins.get('top', 1.0))
                self.doc.sections[0].bottom_margin = Inches(margins.get('bottom', 1.0))
                self.doc.sections[0].left_margin = Inches(margins.get('left', 1.25))
                self.doc.sections[0].right_margin = Inches(margins.get('right', 1.25))
            
            # 设置页面方向
            orientation = page_layout.get('orientation', 'portrait')
            if orientation == 'landscape':
                self.doc.sections[0].orientation = 1  # 横向
            else:
                self.doc.sections[0].orientation = 0  # 纵向
    
    def _update_counters(self, level):
        """更新标题计数器"""
        self.counters[level-1] += 1
        for i in range(level, len(self.counters)):
            self.counters[i] = 0
    
    def _process_heading(self, element, level):
        """处理标题元素（带多级编号）"""
        self._update_counters(level)
        # 获取标题配置，如果不存在则使用默认配置
        heading_config = self.style_config.get('headings', {}).get(f'h{level}', {})
        if not heading_config:
            # 使用默认样式配置
            default_config = self.style_config.get('default_styles', {})
            heading_config = {
                'font': default_config.get('font', {}),
                'paragraph': default_config.get('paragraph', {})
            }
        # 使用内置标题（仅用于层级/样式），编号使用前缀文本实现
        heading = self.doc.add_heading('', level=level)
        
        # 生成前缀文本编号（依据配置的 numbering.format）
        heading_text = element.get_text()
        numbering_cfg = heading_config.get('numbering', {})
        fmt = numbering_cfg.get('format')
        prefix = ''
        if fmt:
            import re as _re
            tokens = _re.findall(r'%([1-6])', fmt)
            if tokens:
                values = []
                for t in tokens:
                    idx = int(t)
                    values.append(str(self.counters[idx]))
                def _rep(m):
                    i = int(m.group(1))
                    return values[tokens.index(str(i))]
                prefix = _re.sub(r'%([1-6])', _rep, fmt)
        else:
            mapped_ilvl = max(0, level - 2)
            if mapped_ilvl >= 0:
                nums = [str(self.counters[i]) for i in range(1, 1 + mapped_ilvl + 1)]
                if nums:
                    prefix = '.'.join(nums)
        if prefix:
            prefix += ' '
        run = heading.add_run(prefix + heading_text)
        
        # 应用字体样式
        font_config = heading_config.get('font', {})
        if font_config:
            # 设置样式字体
            if 'western' in font_config:
                self.doc.styles[f'Heading {level}'].font.name = font_config['western']
                run.font.name = font_config['western']
            if 'east_asian' in font_config:
                self.doc.styles[f'Heading {level}']._element.rPr.rFonts.set(qn('w:eastAsia'), font_config['east_asian'])
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_config['east_asian'])
            if 'size' in font_config:
                self.doc.styles[f'Heading {level}'].font.size = font_config['size']
                run.font.size = font_config['size']
            if 'color' in font_config:
                self.doc.styles[f'Heading {level}'].font.color.rgb = font_config['color']
                run.font.color.rgb = font_config['color']
            if 'bold' in font_config:
                self.doc.styles[f'Heading {level}'].font.bold = font_config['bold']
                run.font.bold = font_config['bold']
            if 'italic' in font_config:
                self.doc.styles[f'Heading {level}'].font.italic = font_config['italic']
                run.font.italic = font_config['italic']
        
        # 段落设置
        para_config = heading_config.get('paragraph', {})
        if para_config:
            if 'alignment' in para_config:
                # 强制设置对齐，覆盖默认样式
                alignment_value = getattr(WD_PARAGRAPH_ALIGNMENT, para_config['alignment'].upper())
                heading.alignment = alignment_value
                # 同时设置段落格式的对齐
                heading.paragraph_format.alignment = alignment_value
                # 关键：同步到样式，避免样式里的默认居中覆盖
                try:
                    self.doc.styles['ProgrammaticNumberingStyle'].paragraph_format.alignment = alignment_value
                except Exception:
                    pass
                # 兜底：直接写入底层XML的对齐设置
                try:
                    pPr = heading._element.get_or_add_pPr()
                    # 移除已有 jc
                    for child in list(pPr):
                        if child.tag == qn('w:jc'):
                            pPr.remove(child)
                    jc = OxmlElement('w:jc')
                    jc.set(qn('w:val'), para_config['alignment'].lower())
                    pPr.append(jc)
                except Exception:
                    pass
            # 段前距离归零
            try:
                heading.paragraph_format.space_before = 0
                heading.paragraph_format.left_indent = 0
                heading.paragraph_format.first_line_indent = 0
            except Exception:
                pass

        # 不注入 numPr，保持无关联编号
            if 'space_before' in para_config:
                heading.space_before = para_config['space_before']
            if 'space_after' in para_config:
                heading.space_after = para_config['space_after']
            if 'line_spacing' in para_config:
                heading.paragraph_format.line_spacing = para_config['line_spacing']
        
    def _process_paragraph(self, element, parent_paragraph=None):
        """处理段落元素"""
        paragraph = parent_paragraph or self.doc.add_paragraph()
        
        # 应用默认段落样式
        default_config = self.style_config.get('default_styles', {})
        if default_config:
            font_config = default_config.get('font', {})
            para_config = default_config.get('paragraph', {})
            
            # 设置段落字体
            if font_config:
                paragraph.style.font.name = font_config.get('western', 'Times New Roman')
                paragraph.style._element.rPr.rFonts.set(qn('w:eastAsia'), font_config.get('east_asian', '宋体'))
                if 'size' in font_config:
                    paragraph.style.font.size = font_config['size']
                if 'color' in font_config:
                    paragraph.style.font.color.rgb = font_config['color']
            
            # 设置段落格式
            if para_config:
                if 'line_spacing' in para_config:
                    paragraph.paragraph_format.line_spacing = para_config['line_spacing']
                if 'space_before' in para_config:
                    paragraph.paragraph_format.space_before = para_config['space_before']
                if 'space_after' in para_config:
                    paragraph.paragraph_format.space_after = para_config['space_after']
        
        for child in element.children:
            self._process_element(child, paragraph)

    # ===== 数学公式（LaTeX -> MathML -> OMML）支持 =====
    def _get_xslt_transformer(self, xsl_path="/home/yzy/document/project/AI_office-main/config/mml2omml.xsl"):
        if self._xslt_transform is None:
            try:
                with open(xsl_path, 'rb') as f:
                    xslt_root = etree.XML(f.read())
                self._xslt_transform = etree.XSLT(xslt_root)
            except Exception:
                self._xslt_transform = None
        return self._xslt_transform

    def _mathml_to_omml(self, mathml_str: str) -> Optional[str]:
        try:
            transformer = self._get_xslt_transformer()
            if transformer is None:
                return None
            mathml_doc = etree.fromstring(mathml_str.encode('utf-8'))
            omml_doc = transformer(mathml_doc)
            return str(omml_doc)
        except Exception:
            return None

    def _latex_to_omml(self, latex_expr: str) -> Optional[str]:
        try:
            mathml = latex_to_mathml(latex_expr)
        except Exception:
            return None
        return self._mathml_to_omml(mathml)

    def _append_omml_to_paragraph(self, paragraph, omml_xml: str, center: bool = False):
        try:
            omml_elem = parse_xml(omml_xml)
            paragraph._p.append(omml_elem)
            if center:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except Exception:
            pass

    def _append_text_with_math(self, paragraph, text: str):
        import re
        if not text:
            return
        pattern = re.compile(r"(\$\$(.+?)\$\$)|(\$(.+?)\$)", re.DOTALL)
        last_idx = 0
        for m in pattern.finditer(text):
            # 先添加匹配前的普通文本
            if m.start() > last_idx:
                paragraph.add_run(text[last_idx:m.start()])
            expr = None
            is_display = False
            if m.group(1):
                expr = m.group(2)
                is_display = True
            else:
                expr = m.group(4)
                is_display = False
            if expr is not None:
                omml_xml = self._latex_to_omml(expr.strip())
                if omml_xml:
                    if is_display:
                        # 独立段落，居中显示
                        display_para = self.doc.add_paragraph()
                        self._append_omml_to_paragraph(display_para, omml_xml, center=True)
                    else:
                        # 行内插入到当前段落
                        self._append_omml_to_paragraph(paragraph, omml_xml, center=False)
                else:
                    # 失败则原样插入
                    paragraph.add_run(m.group(0))
            last_idx = m.end()
        # 追加剩余文本
        if last_idx < len(text):
            paragraph.add_run(text[last_idx:])
    
    def _process_formatting(self, element, parent_paragraph, style):
        """处理文本格式（加粗/斜体）"""
        run = parent_paragraph.add_run(element.get_text())
        
        # 获取文本样式配置
        text_styles = self.style_config.get('text_styles', {})
        style_config = text_styles.get(style, {})
        
        if style == 'bold':
            run.bold = True
            # 应用粗体样式配置
            if 'font' in style_config:
                font_config = style_config['font']
                if 'size' in font_config:
                    run.font.size = font_config['size']
                if 'color' in font_config:
                    run.font.color.rgb = font_config['color']
        elif style == 'italic':
            run.italic = True
            # 应用斜体样式配置
            if 'font' in style_config:
                font_config = style_config['font']
                if 'size' in font_config:
                    run.font.size = font_config['size']
                if 'color' in font_config:
                    run.font.color.rgb = font_config['color']
        elif style == 'code':
            # 应用代码样式配置
            if 'font' in style_config:
                font_config = style_config['font']
                if 'western' in font_config:
                    run.font.name = font_config['western']
                if 'size' in font_config:
                    run.font.size = font_config['size']
                if 'color' in font_config:
                    run.font.color.rgb = font_config['color']
            # 应用段落背景色
            if 'paragraph' in style_config and 'background' in style_config['paragraph']:
                # 这里需要设置段落背景色，但python-docx对背景色支持有限
                pass
    
    def _process_list(self, element, list_type):
        """处理列表元素"""
        # 获取列表样式配置
        lists_config = self.style_config.get('lists', {})
        if list_type == 'ul':
            list_config = lists_config.get('unordered', {})
            list_style = list_config.get('style', 'List Bullet')
        else:
            list_config = lists_config.get('ordered', {})
            list_style = list_config.get('style', 'List Number')
        
        for li in element.find_all('li', recursive=False):
            paragraph = self.doc.add_paragraph(style=list_style)
            
            # 应用列表字体样式
            if 'font' in list_config:
                font_config = list_config['font']
                if 'size' in font_config:
                    paragraph.style.font.size = font_config['size']
            
            for child in li.children:
                self._process_element(child, paragraph)
    
    def _resolve_image_path(self, src: str):
        if os.path.isabs(src):
            return src
        if hasattr(self, '_input_base_dir'):
            return os.path.join(self._input_base_dir, src)
        return src

    

    def _process_image(self, img_tag: Tag, caption_text: str = None):
        """处理图片 <img>（支持对齐、尺寸、间距、标题）"""
        src = img_tag.get('src')
        if not src:
            return
        image_path = self._resolve_image_path(src)

        # 读取配置
        images_cfg = (self.style_config.get('images') or {}).get('default', {})
        align = (images_cfg.get('alignment') or 'center').lower()
        cfg_width = images_cfg.get('width')
        cfg_height = images_cfg.get('height')
        space_before = images_cfg.get('space_before')
        space_after = images_cfg.get('space_after')

        # 检查 <img> 自身的宽高（像素）并转换为英寸（假设 96 DPI）
        width_in = None
        height_in = None
        try:
            if img_tag.has_attr('width'):
                w_attr = img_tag.get('width')
                if isinstance(w_attr, str) and w_attr.endswith('px'):
                    width_in = float(w_attr[:-2]) / 96.0
                else:
                    width_in = float(w_attr) / 96.0
        except Exception:
            width_in = None
        try:
            if img_tag.has_attr('height'):
                h_attr = img_tag.get('height')
                if isinstance(h_attr, str) and h_attr.endswith('px'):
                    height_in = float(h_attr[:-2]) / 96.0
                else:
                    height_in = float(h_attr) / 96.0
        except Exception:
            height_in = None

        # 若标签未给定，则使用配置
        if width_in is None and isinstance(cfg_width, (int, float)):
            width_in = cfg_width
        if height_in is None and isinstance(cfg_height, (int, float)):
            height_in = cfg_height

        # 插入图片（使用文档级 add_picture 生成独立段落，便于设置对齐与间距）
        try:
            if width_in is not None and (height_in is None or height_in == 'auto'):
                self.doc.add_picture(image_path, width=Inches(width_in))
            elif height_in is not None and (width_in is None or width_in == 'auto'):
                self.doc.add_picture(image_path, height=Inches(height_in))
            elif width_in is not None and height_in is not None and isinstance(width_in, (int, float)) and isinstance(height_in, (int, float)):
                # 优先使用宽度，保持比例
                self.doc.add_picture(image_path, width=Inches(width_in))
            else:
                self.doc.add_picture(image_path)
        except Exception:
            # 插入失败则跳过
            return

        # 设置图片所在段落的对齐与间距
        if self.doc.paragraphs:
            pic_para = self.doc.paragraphs[-1]
            if align == 'left':
                pic_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif align == 'right':
                pic_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:
                pic_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            try:
                if space_before is not None:
                    pic_para.paragraph_format.space_before = space_before
                if space_after is not None:
                    pic_para.paragraph_format.space_after = space_after
            except Exception:
                pass

        # 生成标题（优先 figcaption 传入的 caption_text，其次使用 alt）
        final_caption = caption_text or img_tag.get('alt')
        if final_caption:
            cap_cfg = (images_cfg.get('caption_style') or {})
            cap_para = self.doc.add_paragraph()
            cap_run = cap_para.add_run(str(final_caption))

            # 字体样式
            font_cfg = (cap_cfg.get('font') or {})
            try:
                if 'size' in font_cfg:
                    cap_run.font.size = font_cfg['size']
                if 'color' in font_cfg:
                    cap_run.font.color.rgb = font_cfg['color']
                if font_cfg.get('italic') is True:
                    cap_run.italic = True
                if font_cfg.get('bold') is True:
                    cap_run.bold = True
            except Exception:
                pass

            # 段落样式
            para_cfg = (cap_cfg.get('paragraph') or {})
            try:
                align_str = (para_cfg.get('alignment') or align).upper()
                cap_para.alignment = getattr(WD_PARAGRAPH_ALIGNMENT, align_str)
            except Exception:
                pass
            try:
                if 'space_before' in para_cfg:
                    cap_para.paragraph_format.space_before = para_cfg['space_before']
                if 'space_after' in para_cfg:
                    cap_para.paragraph_format.space_after = para_cfg['space_after']
            except Exception:
                pass

    def _process_table(self, element: Tag):
        """处理表格元素 <table>（含 thead/tbody/thead/th/td）"""
        tables_cfg = (self.style_config.get('tables') or {}).get('default', {})
        style_name = tables_cfg.get('style', 'Table Grid')
        align = (tables_cfg.get('alignment') or 'center').lower()

        # 收集行
        rows = []
        for tr in element.find_all('tr', recursive=True):
            cells = []
            for cell in tr.find_all(['th', 'td'], recursive=False):
                # 收集纯文本（保留子节点文本）
                cells.append(cell.get_text(strip=True))
            if cells:
                rows.append(cells)
        if not rows:
            return

        cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=cols)
        try:
            table.style = style_name
        except Exception:
            pass

        # 对齐
        if align == 'left':
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
        elif align == 'right':
            table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        else:
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 填充内容
        for r_idx, row in enumerate(rows):
            for c_idx in range(cols):
                text = row[c_idx] if c_idx < len(row) else ''
                cell = table.cell(r_idx, c_idx)
                cell.text = text

        # 列宽
        col_cfg = (tables_cfg.get('column_widths') or {})
        default_w = col_cfg.get('default', None)
        unit = (col_cfg.get('unit') or 'inches').lower()
        if isinstance(default_w, (int, float)):
            width_length = Inches(default_w) if unit == 'inches' else None
            if width_length is not None:
                try:
                    for c_idx in range(cols):
                        for r_idx in range(len(rows)):
                            table.cell(r_idx, c_idx).width = width_length
                except Exception:
                    pass

    def _process_element(self, element, parent_paragraph=None):
        """递归处理HTML元素"""
        if isinstance(element, Tag):
            # 跳过已处理的节点（如图片标题）
            if element.has_attr('data-processed'):
                return
            # 设计约定：# 作为文章标题，不参与编号；从 ## 开始并保持原生级别映射
            # 即：## -> h2，### -> h3
            if element.name.startswith('h') and len(element.name) == 2 and int(element.name[1]) > 1:
                level = int(element.name[1])
                self._process_heading(element, level)
            elif element.name == 'p':
                self._process_paragraph(element, parent_paragraph)
            elif element.name == 'img':
                # 单独处理图片为独立段落
                self._process_image(element)
                element['data-processed'] = 'true'
                return
            elif element.name == 'figure':
                # figure 中包含 img 与可选 figcaption
                img = element.find('img')
                if img is not None:
                    caption = None
                    cap_tag = element.find('figcaption')
                    if cap_tag is not None:
                        caption = cap_tag.get_text(strip=True)
                        cap_tag['data-processed'] = 'true'
                    self._process_image(img, caption_text=caption)
                    element['data-processed'] = 'true'
                    return
            elif element.name in ['strong', 'b']:
                self._process_formatting(element, parent_paragraph, 'bold')
            elif element.name in ['em', 'i']:
                self._process_formatting(element, parent_paragraph, 'italic')
            elif element.name == 'br' and parent_paragraph:
                parent_paragraph.add_run('\n')
            elif element.name in ['ul', 'ol']:
                self._process_list(element, element.name)
            elif element.name in ['code', 'pre']:
                self._process_formatting(element, parent_paragraph, 'code')
            elif element.name == 'table':
                self._process_table(element)
            else:
                for child in element.children:
                    self._process_element(child, parent_paragraph)
        elif parent_paragraph is not None:
            # 处理文本节点（支持 $...$ 与 $$...$$ 数学公式）
            self._append_text_with_math(parent_paragraph, str(element))
    
    def convert(self, input_file, output_file):
        """执行转换过程"""
        # 读取Markdown文件
        with open(input_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 转换Markdown为HTML
        md = Markdown(extensions=['toc', 'tables'])
        html_content = md.convert(md_content)
        
        # 创建文档
        self.doc = self._create_document()
        # 记录输入文件所在目录，供相对路径资源（图片等）解析
        self._input_base_dir = os.path.dirname(os.path.abspath(input_file))
        
        # 应用页面布局设置
        self._apply_page_layout()
        
        # 解析HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 处理所有子元素
        for child in soup.children:
            self._process_element(child)
        
        # 保存文档
        self.doc.save(output_file)
        print(f"Successfully converted {input_file} to {output_file}")

# 使用示例
if __name__ == "__main__":
    converter = MarkdownToDocxConverter()
    converter.convert('./md_input_file/input.md', './output_file/output.docx')
